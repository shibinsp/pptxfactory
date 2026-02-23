"""
Document Processor - Handles PDF, Word, Image uploads
Extracts content and converts to presentation format
"""

import os
import io
import re
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
import base64


class DocumentProcessor:
    """
    Process uploaded documents and extract content for presentations
    Supports: PDF, DOCX, TXT, Images
    """
    
    def __init__(self, mistral_client=None):
        self.mistral_client = mistral_client
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'doc': self._process_docx,
            'txt': self._process_txt,
            'md': self._process_txt,
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image,
            'webp': self._process_image,
            'gif': self._process_image
        }
    
    def process_file(self, file_path: str, file_type: str = None) -> Dict[str, Any]:
        """
        Process uploaded file and extract content
        
        Returns:
            {
                "success": bool,
                "content": str,  # Extracted text
                "slides": List[Dict],  # Suggested slides
                "images": List[str],  # Extracted image paths
                "metadata": Dict,  # File metadata
                "summary": str  # AI-generated summary
            }
        """
        # Detect file type if not provided
        if not file_type:
            file_type = self._detect_file_type(file_path)
        
        file_type = file_type.lower().lstrip('.')
        
        if file_type not in self.supported_formats:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}",
                "supported_formats": list(self.supported_formats.keys())
            }
        
        try:
            processor = self.supported_formats[file_type]
            result = processor(file_path)
            
            # Generate AI summary if Mistral is available
            if self.mistral_client and result.get('content'):
                result['summary'] = self._generate_summary(result['content'])
                result['slides'] = self._suggest_slides(result['content'])
            
            result['success'] = True
            return result
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type from extension"""
        _, ext = os.path.splitext(file_path)
        return ext.lower().lstrip('.')
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF file"""
        try:
            import PyPDF2
            
            content = []
            images = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"--- Page {page_num + 1} ---\n{text}")
                
                metadata = {
                    "pages": num_pages,
                    "title": pdf_reader.metadata.get('/Title', 'Unknown'),
                    "author": pdf_reader.metadata.get('/Author', 'Unknown')
                }
            
            return {
                "content": "\n\n".join(content),
                "images": images,
                "metadata": metadata,
                "file_type": "pdf"
            }
            
        except ImportError:
            return self._process_pdf_alternative(file_path)
    
    def _process_pdf_alternative(self, file_path: str) -> Dict:
        """Alternative PDF processing using pdfplumber"""
        try:
            import pdfplumber
            
            content = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        content.append(f"--- Page {i + 1} ---\n{text}")
                
                metadata = {
                    "pages": len(pdf.pages),
                    "file_type": "pdf"
                }
            
            return {
                "content": "\n\n".join(content),
                "images": [],
                "metadata": metadata,
                "file_type": "pdf"
            }
            
        except ImportError:
            raise Exception("PDF processing requires PyPDF2 or pdfplumber. Install with: pip install PyPDF2 pdfplumber")
    
    def _process_docx(self, file_path: str) -> Dict:
        """Process Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    content.append("\n[TABLE]\n" + "\n".join(table_data))
            
            # Extract images
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append(rel.target_ref)
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "images": len(images)
            }
            
            return {
                "content": "\n\n".join(content),
                "images": images,
                "metadata": metadata,
                "file_type": "docx"
            }
            
        except ImportError:
            raise Exception("DOCX processing requires python-docx. Install with: pip install python-docx")
    
    def _process_txt(self, file_path: str) -> Dict:
        """Process text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
        
        # Try to detect markdown headings
        headings = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
        
        metadata = {
            "lines": len(content.split('\n')),
            "words": len(content.split()),
            "headings": len(headings)
        }
        
        return {
            "content": content,
            "images": [],
            "metadata": metadata,
            "file_type": "txt",
            "headings": headings
        }
    
    def _process_image(self, file_path: str) -> Dict:
        """Process image file - OCR extraction"""
        try:
            from PIL import Image
            
            # Get image info
            with Image.open(file_path) as img:
                width, height = img.size
                format_type = img.format
                mode = img.mode
            
            # Try OCR if available
            ocr_text = ""
            try:
                import pytesseract
                ocr_text = pytesseract.image_to_string(Image.open(file_path))
            except:
                pass
            
            metadata = {
                "width": width,
                "height": height,
                "format": format_type,
                "mode": mode
            }
            
            return {
                "content": ocr_text if ocr_text else "Image file uploaded",
                "images": [file_path],
                "metadata": metadata,
                "file_type": "image",
                "ocr_text": ocr_text
            }
            
        except ImportError:
            return {
                "content": "Image file uploaded",
                "images": [file_path],
                "metadata": {"file_type": "image"},
                "file_type": "image"
            }
    
    def _generate_summary(self, content: str) -> str:
        """Generate AI summary of content"""
        if not self.mistral_client:
            return "Summary not available - AI client not configured"
        
        try:
            # Truncate content if too long
            max_length = 4000
            truncated = content[:max_length] + "..." if len(content) > max_length else content
            
            prompt = f"""Summarize the following document in 3-5 bullet points. 
            Focus on the main topics and key takeaways:
            
            {truncated}
            
            Summary:"""
            
            response = self.mistral_client.chat.complete(
                model="mistral-tiny",
                messages=[{"role": "user", "content": prompt}]
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            return f"Error generating summary: {str(e)}"
    
    def _suggest_slides(self, content: str) -> List[Dict]:
        """Suggest slide structure from content"""
        if not self.mistral_client:
            return self._basic_slide_suggestion(content)
        
        try:
            max_length = 3000
            truncated = content[:max_length] + "..." if len(content) > max_length else content
            
            prompt = f"""Based on this document, suggest a presentation structure.
            Return ONLY a JSON array with slide objects containing 'title' and 'content'.
            
            Document:
            {truncated}
            
            Suggested slides (JSON format):
            """
            
            response = self.mistral_client.chat.complete(
                model="mistral-tiny",
                messages=[{"role": "user", "content": prompt}]
            )
            
            # Extract JSON from response
            response_text = response.choices[0].message.content
            
            # Try to find JSON array in response
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                slides = json.loads(json_match.group())
                return slides
            
            return self._basic_slide_suggestion(content)
            
        except Exception as e:
            return self._basic_slide_suggestion(content)
    
    def _basic_slide_suggestion(self, content: str) -> List[Dict]:
        """Basic slide suggestion without AI"""
        # Split by headings or paragraphs
        sections = re.split(r'\n\n+', content)
        
        slides = []
        
        # Title slide
        if sections:
            first_lines = sections[0].split('\n')[:3]
            title = first_lines[0][:100] if first_lines else "Presentation"
            subtitle = first_lines[1][:200] if len(first_lines) > 1 else ""
            
            slides.append({
                "type": "title",
                "title": title,
                "content": subtitle
            })
        
        # Content slides - group sections
        for i, section in enumerate(sections[1:6], 1):  # Max 5 content slides
            lines = section.split('\n')
            title = lines[0][:100] if lines else f"Slide {i+1}"
            content_text = '\n'.join(lines[1:5]) if len(lines) > 1 else section[:300]
            
            # Convert to bullet points
            bullets = [line.strip() for line in content_text.split('\n') if line.strip()][:5]
            
            slides.append({
                "type": "content",
                "title": title,
                "content": bullets
            })
        
        # Closing slide
        slides.append({
            "type": "closing",
            "title": "Thank You",
            "content": "Questions?"
        })
        
        return slides
    
    def extract_topics(self, content: str) -> List[str]:
        """Extract main topics from content"""
        # Simple keyword extraction
        words = re.findall(r'\b[A-Z][a-z]+\b', content)
        word_freq = {}
        
        for word in words:
            if len(word) > 3:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Return top topics
        sorted_topics = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [topic for topic, freq in sorted_topics[:10]]
    
    def convert_to_presentation_format(self, processed_doc: Dict) -> Dict:
        """Convert processed document to presentation format"""
        slides = processed_doc.get('slides', [])
        
        presentation = {
            "title": "Document Presentation",
            "slides": []
        }
        
        for i, slide_data in enumerate(slides):
            slide = {
                "id": f"slide-{i+1}",
                "order": i,
                "layout": slide_data.get('type', 'content'),
                "title": slide_data.get('title', f"Slide {i+1}"),
                "content": slide_data.get('content', '')
            }
            presentation["slides"].append(slide)
        
        return presentation


# Utility functions for file handling
def get_file_extension(filename: str) -> str:
    """Get file extension from filename"""
    return os.path.splitext(filename)[1].lower().lstrip('.')


def is_supported_format(filename: str) -> bool:
    """Check if file format is supported"""
    ext = get_file_extension(filename)
    supported = ['pdf', 'docx', 'doc', 'txt', 'md', 'png', 'jpg', 'jpeg', 'webp', 'gif']
    return ext in supported


def get_file_icon(filename: str) -> str:
    """Get appropriate icon for file type"""
    ext = get_file_extension(filename)
    
    icons = {
        'pdf': '📄',
        'docx': '📝',
        'doc': '📝',
        'txt': '📃',
        'md': '📑',
        'png': '🖼️',
        'jpg': '🖼️',
        'jpeg': '🖼️',
        'webp': '🖼️',
        'gif': '🎬'
    }
    
    return icons.get(ext, '📎')
